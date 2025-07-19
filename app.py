import os
from docx import Document
from utils import extract_text_from_pdf
from reasoner import generate_reasoning

# 📥 1. Get job title and description from user
job_title = input("Enter the job title: ")
job_description = input("Enter the full job description: ")

# 📁 2. Resumes folder path
RESUME_FOLDER = "data/resumes"
OUTPUT_FOLDER = "results"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# 📂 3. Process all PDFs in the folder
for filename in os.listdir(RESUME_FOLDER):
    if filename.lower().endswith(".pdf"):
        resume_path = os.path.join(RESUME_FOLDER, filename)
        print(f"\n📄 Parsing {filename}...")

        try:
            # 🔍 Extract resume text
            resume_text = extract_text_from_pdf(resume_path)

            # 🧠 Generate reasoning with model
            print("🤖 Generating reasoning using Mistral 7B...")
            result = generate_reasoning(job_description, resume_text)
            print("✅ Match reasoning generated.\n")

            # 📄 Prepare output filename
            base_name = os.path.splitext(filename)[0]
            output_path = os.path.join(OUTPUT_FOLDER, f"{base_name}_result.docx")

            # 📝 Write to Word file
            doc = Document()
            doc.add_heading(f"Match Result for {base_name}", 0)
            doc.add_paragraph(f"Job Title: {job_title}")
            doc.add_paragraph("Job Description:\n" + job_description)
            doc.add_paragraph("Resume Filename: " + filename)
            doc.add_paragraph("\nReasoning:\n" + result)

            doc.save(output_path)
            print(f"💾 Saved result to: {output_path}")

        except Exception as e:
            print(f"❌ Error processing {filename}: {e}")
